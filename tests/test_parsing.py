"""Generated fixtures stay in pytest's temp directory; no production data/API."""
import importlib.util
from pathlib import Path
import sys

import pytest

ROOT = Path(__file__).resolve().parents[1]
spec = importlib.util.spec_from_file_location("kb_parsing", ROOT / "services/worker/app/parsing.py")
parsing = importlib.util.module_from_spec(spec)
sys.modules[spec.name] = parsing
spec.loader.exec_module(parsing)
Block, PdfPage = parsing.Block, parsing.PdfPage


def line(text, top, bottom=None, **extra):
    return dict(kind="text", text=text, top=top, bottom=bottom or top + 12, x0=50, x1=550, **extra)


def table(rows, top, bottom):
    return dict(kind="table", rows=rows, top=top, bottom=bottom, x0=50, x1=550)


def test_pdf_continuation_and_margins():
    pages = [PdfPage(1, 600, 800, [line("企业制度", 20), line("请假申请需要", 690), line("第1页", 765)]),
             PdfPage(2, 600, 800, [line("企业制度", 20), line("直属主管审批。", 85), line("第2页", 765)])]
    blocks = parsing.assemble_pdf(pages)
    assert len(blocks) == 1
    assert blocks[0].text == "请假申请需要直属主管审批。"
    chunks = parsing.split_blocks(blocks)
    assert (chunks[0].page_start, chunks[0].page_end) == (1, 2)
    assert chunks[0].metadata["cross_page_continuation"]


@pytest.mark.parametrize("next_text,last_text", [("第二章 劳动纪律", "请假申请需要"), ("新的内容", "条款已结束。"),
                                                    ("（一）下一条", "请假申请需要")])
def test_pdf_does_not_join_new_sections(next_text, last_text):
    pages = [PdfPage(1, 600, 800, [line(last_text, 690)]), PdfPage(2, 600, 800, [line(next_text, 85)])]
    blocks = parsing.assemble_pdf(pages)
    assert len(blocks) == 2
    assert not any(b.metadata.get("cross_page_continuation") for b in blocks)


def test_pdf_preserves_body_repetition_and_unique_margin_text():
    pages = [PdfPage(n, 600, 800, [line(f"生效年份{2020+n}", 20), line("重要重复条款", 300)]) for n in (1, 2, 3)]
    parsing.clean_pdf_margins(pages)
    assert all(len(p.elements) == 2 for p in pages)


def test_pdf_does_not_join_over_blank_page():
    blocks = parsing.assemble_pdf([PdfPage(1, 600, 800, [line("尚未结束", 690)]), PdfPage(2, 600, 800, []),
                                   PdfPage(3, 600, 800, [line("不是续段", 85)])])
    assert len(blocks) == 2


def test_pdf_continued_table_carries_header_and_pages():
    pages = [PdfPage(1, 600, 800, [table([["假别", "天数"], ["事假", "1"]], 600, 715)]),
             PdfPage(2, 600, 800, [table([["假别", "天数"], ["年假", "5"]], 85, 200)])]
    blocks = parsing.assemble_pdf(pages)
    assert [b.metadata["table_id"] for b in blocks] == [1, 1]
    assert blocks[1].metadata["continued_table"]
    assert "假别：年假" in blocks[1].text
    chunk = parsing.split_blocks(blocks)[0]
    assert (chunk.page_start, chunk.page_end) == (1, 2)
    assert "表头：假别 | 天数" in chunk.text
    assert len(chunk.metadata["sources"]) == 2


def test_pdf_different_table_is_not_continuation():
    pages = [PdfPage(1, 600, 800, [table([["假别", "天数"], ["事假", "1"]], 600, 715)]),
             PdfPage(2, 600, 800, [table([["岗位", "等级"], ["技术员", "5"]], 85, 200)])]
    blocks = parsing.assemble_pdf(pages)
    assert [b.metadata["table_id"] for b in blocks] == [1, 2]


def test_pdf_same_table_away_from_page_boundary_is_separate():
    rows = [["假别", "天数"], ["事假", "1"]]
    blocks = parsing.assemble_pdf([PdfPage(1, 600, 800, [table(rows, 100, 200)]),
                                   PdfPage(2, 600, 800, [table(rows, 85, 200)])])
    assert blocks[1].metadata["table_id"] == 2


def test_docx_interleaves_tables_and_preserves_headings(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("请假管理", level=1)
    doc.add_paragraph("表格之前的申请条件")
    grid = doc.add_table(rows=2, cols=2)
    for cell, value in zip([c for row in grid.rows for c in row.cells], ["假别", "天数", "年假", "5"]):
        cell.text = value
    doc.add_paragraph("表格之后的例外说明")
    doc.add_heading("考勤管理", level=1)
    doc.add_paragraph("下一章内容")
    path = tmp_path / "policy.docx"
    doc.save(path)
    blocks = parsing.extract(path, path.name)
    joined = "\n".join(b.text for b in blocks)
    assert joined.index("申请条件") < joined.index("假别：年假") < joined.index("例外说明") < joined.index("考勤管理")
    assert next(b for b in blocks if b.kind == "table").metadata["headings"] == ["请假管理"]
    chunks = parsing.split_blocks(blocks)
    assert all(c.page_start is None for c in chunks)
    assert not any("下一章内容" in c.text and "申请条件" in c.text for c in chunks)


def test_excel_every_chunk_has_sheet_headers_and_row_metadata(tmp_path):
    from openpyxl import Workbook
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "人资台账"
    sheet.append(["姓名", "部门", "备注"])
    for n in range(30):
        sheet.append([f"员工{n}", "人力资源部", "待确认" * 12])
    other = workbook.create_sheet("技术台账")
    other.append(["图号", "库存"])
    other.append(["A-001", 0])
    path = tmp_path / "tables.xlsx"
    workbook.save(path)
    blocks = parsing.extract(path, path.name)
    chunks = parsing.split_blocks(blocks, size=220, overlap=25)
    assert len(chunks) > 5
    assert all("工作表：" in c.text and "表头：" in c.text for c in chunks)
    assert all(c.metadata.get("sheet") and c.metadata.get("row_start") for c in chunks)
    assert all(not ("人资台账" in c.text and "技术台账" in c.text) for c in chunks)
    assert "库存：0" in chunks[-1].text
    assert all(len(c.text) <= 220 for c in chunks)


def test_excel_merged_multilevel_header_and_title(tmp_path):
    from openpyxl import Workbook
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["员工信息表"])
    sheet.merge_cells("A1:C1")
    sheet.append(["员工", "收入", None])
    sheet.merge_cells("B2:C2")
    sheet.append(["姓名", "基本工资", "补贴"])
    sheet.append(["张三", 100, 0])
    path = tmp_path / "headers.xlsx"
    workbook.save(path)
    blocks = parsing.extract(path, path.name)
    assert blocks[0].metadata["headers"] == ["员工 / 姓名", "收入 / 基本工资", "收入 / 补贴"]
    assert "员工信息表" in blocks[0].context
    assert "收入 / 补贴：0" in blocks[0].text
    assert blocks[0].metadata["row_start"] == 4


def test_excel_headerless_numbers_not_discarded(tmp_path):
    from openpyxl import Workbook
    workbook = Workbook()
    workbook.active.append([1, 2])
    workbook.active.append([3, 4])
    path = tmp_path / "numbers.xlsx"
    workbook.save(path)
    blocks = parsing.extract(path, path.name)
    assert "列1：1" in blocks[0].text
    assert len(blocks) == 2


def test_long_line_strict_limit_and_no_content_loss():
    text = "".join(chr(0x4e00 + n) for n in range(2000))
    chunks = parsing.split_blocks([Block(text, page_start=3, page_end=4)], size=160, overlap=20)
    assert all(len(c.text) <= 160 for c in chunks)
    assert all((c.page_start, c.page_end) == (3, 4) for c in chunks)
    assert set(text) == set("".join(c.text for c in chunks))


def test_long_table_row_repeats_context():
    blocks = parsing.table_blocks([(1, ["姓名", "说明"]), (2, ["张三", "长说明" * 1000])],
                                  context="工作表：登记表", metadata={"sheet": "登记表"})
    chunks = parsing.split_blocks(blocks, size=200, overlap=20)
    assert len(chunks) > 10
    assert all(c.text.startswith("工作表：登记表\n表头：姓名 | 说明\n") for c in chunks)


def test_oversized_header_fails_explicitly():
    with pytest.raises(ValueError, match="表头或标题"):
        parsing.split_blocks([Block("value", context="x" * 200)], size=200)


@pytest.mark.parametrize("size,overlap", [(10, 1), (100, -1), (100, 100)])
def test_invalid_chunk_parameters(size, overlap):
    with pytest.raises(ValueError):
        parsing.split_blocks([Block("text")], size=size, overlap=overlap)


def test_plain_text_dispatch(tmp_path):
    path = tmp_path / "content.txt"
    path.write_text("原始内容\n第二段", encoding="utf-8")
    assert parsing.extract(path, path.name)[0].text == "原始内容\n第二段"
    with pytest.raises(ValueError, match="不支持"):
        parsing.extract(path, "file.dwg")


def write_pdf(path, streams, *, image_page=None):
    """Minimal actual PDF fixture, using built-in font and optional scan image."""
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, NumberObject, DecodedStreamObject
    writer = PdfWriter()
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"),
                             NameObject("/BaseFont"): NameObject("/Helvetica")})
    for number, stream in enumerate(streams, 1):
        page = writer.add_blank_page(600, 800)
        resources = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})})
        if number == image_page:
            img = DecodedStreamObject()
            img.set_data(b"\xff\xff\xff" * 100)
            img.update({NameObject("/Type"): NameObject("/XObject"), NameObject("/Subtype"): NameObject("/Image"),
                        NameObject("/Width"): NumberObject(10), NameObject("/Height"): NumberObject(10),
                        NameObject("/ColorSpace"): NameObject("/DeviceRGB"), NameObject("/BitsPerComponent"): NumberObject(8)})
            resources[NameObject("/XObject")] = DictionaryObject({NameObject("/Im0"): writer._add_object(img)})
            stream += " q 500 0 0 600 50 100 cm /Im0 Do Q"
        page[NameObject("/Resources")] = resources
        contents = DecodedStreamObject()
        contents.set_data(stream.encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(contents)
    with path.open("wb") as output:
        writer.write(output)


def pdf_text(text, x, y):
    return f"BT /F1 12 Tf {x} {y} Td ({text}) Tj ET\n"


def test_real_pdf_layout_extraction(tmp_path):
    path = tmp_path / "pages.pdf"
    write_pdf(path, [pdf_text("Company policy", 50, 780) + pdf_text("Application requires", 50, 100) + pdf_text("1", 300, 20),
                     pdf_text("Company policy", 50, 780) + pdf_text("manager approval.", 50, 700) + pdf_text("2", 300, 20)])
    chunks = parsing.split_blocks(parsing.extract(path, path.name))
    assert len(chunks) == 1
    assert "Application requires manager approval." in chunks[0].text
    assert "Company policy" not in chunks[0].text
    assert (chunks[0].page_start, chunks[0].page_end) == (1, 2)


def test_real_pdf_table_extraction_preserves_order(tmp_path):
    path = tmp_path / "table.pdf"
    grid = "0.5 w " + " ".join(f"{x} 500 m {x} 580 l S" for x in (50, 250, 450))
    grid += " " + " ".join(f"50 {y} m 450 {y} l S" for y in (500, 540, 580))
    stream = pdf_text("Before the table.", 50, 650) + grid
    stream += "\n" + pdf_text("Type", 60, 555) + pdf_text("Days", 260, 555)
    stream += pdf_text("Annual", 60, 515) + pdf_text("5", 260, 515) + pdf_text("After the table.", 50, 400)
    write_pdf(path, [stream])
    blocks = parsing.extract(path, path.name)
    assert [b.kind for b in blocks] == ["text", "table", "text"]
    assert "Type：Annual" in blocks[1].text
    assert "Days：5" in blocks[1].text


def test_mixed_pdf_ocr_only_scanned_page(tmp_path, monkeypatch):
    path = tmp_path / "mixed.pdf"
    write_pdf(path, [pdf_text("A digital page with sufficient embedded text.", 50, 650),
                     pdf_text("Long native header on a page with a scanned body.", 50, 780)], image_page=2)
    calls = []
    def fake_ocr(path, number):
        calls.append(number)
        return PdfPage(number, 600, 800, [line("OCR scan text.", 200, extraction="ocr")])
    monkeypatch.setattr(parsing, "_ocr_pdf_page", fake_ocr)
    blocks = parsing.extract(path, path.name)
    assert calls == [2]
    assert blocks[-1].metadata["extraction"] == "ocr"
    assert "OCR scan text." in blocks[-1].text


def test_worker_persists_page_span_and_metadata(monkeypatch):
    # Import through its real package, but stub every side-effecting boundary.
    import json
    from services.worker.app import main
    calls = []
    class FakeCursor:
        lastrowid = 123
        def __enter__(self): return self
        def __exit__(self, *args): pass
        def execute(self, sql, params): calls.append((sql, params))
    class FakeConnection:
        def __enter__(self): return self
        def __exit__(self, *args): pass
        def cursor(self): return FakeCursor()
        def commit(self): pass
    monkeypatch.setattr(main, "db", FakeConnection)
    monkeypatch.setattr(main, "cleanup_indexes", lambda _: None)
    chunks = [parsing.Chunk("跨页内容", 2, 3, {"source": "worker", "sheet": "测试", "parser_version": parsing.PARSER_VERSION})]
    units = main.save_units({"document_id": 10, "document_version_id": 20}, chunks)
    params = calls[-1][1]
    assert params[2:4] == (2, 3)
    assert json.loads(params[-1])["parser_version"] == parsing.PARSER_VERSION
    assert units[0]["id"] == 123


def test_empty_extraction_does_not_delete_existing_indexes(monkeypatch, tmp_path):
    from services.worker.app import main
    path = tmp_path / "empty"
    path.touch()
    monkeypatch.setattr(main, "download", lambda _: path)
    monkeypatch.setattr(main, "extract", lambda *_: [])
    monkeypatch.setattr(main, "save_units", lambda *_: pytest.fail("must not delete existing chunks"))
    with pytest.raises(ValueError, match="未提取"):
        main.run({"job_type": "reindex", "original_filename": "empty.pdf"})
    assert not path.exists()


def test_three_page_table_continuation_keeps_row_numbers():
    pages = [PdfPage(n, 600, 800, [table([["名称", "数量"], [f"物料{n}", str(n)]], 85, 715)]) for n in (1, 2, 3)]
    blocks = parsing.assemble_pdf(pages)
    assert [b.metadata["row_start"] for b in blocks] == [2, 3, 4]
    chunk = parsing.split_blocks(blocks)[0]
    assert (chunk.page_start, chunk.page_end) == (1, 3)


def test_table_header_only_retained():
    blocks = parsing.table_blocks([(1, ["名称", "数量"])], context="空表", metadata={})
    assert "名称" in blocks[0].context
    assert "无数据行" in blocks[0].text


def test_excel_blank_line_separates_tables(tmp_path):
    from openpyxl import Workbook
    workbook = Workbook()
    for row in [["名称", "数量"], ["零件", 3], [], ["员工", "部门"], ["张三", "技术部"]]:
        workbook.active.append(row)
    path = tmp_path / "regions.xlsx"
    workbook.save(path)
    blocks = parsing.extract(path, path.name)
    assert [b.metadata["table_id"] for b in blocks] == [1, 2]
    assert "员工：张三" in blocks[1].text


def test_excel_worker_extensionless_download(tmp_path):
    from openpyxl import Workbook
    workbook = Workbook()
    workbook.active.append(["项目", "状态"])
    workbook.active.append(["审批", "通过"])
    path = tmp_path / "kb-source-no-extension"
    workbook.save(path)
    blocks = parsing.extract(path, "用户上传的台账.xlsx")
    assert "项目：审批" in blocks[0].text
