import hashlib
import logging
import math
import os
import re
import tempfile
import time
from pathlib import Path
from urllib.parse import unquote, urlparse

import httpx
import pymysql
import pytesseract
from docx import Document as DocxDocument
from minio import Minio
from openpyxl import load_workbook
from pdf2image import convert_from_path
from pypdf import PdfReader
from pymilvus import Collection, CollectionSchema, DataType, FieldSchema, connections, utility

logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))
log = logging.getLogger("kb-worker")
COLLECTION = os.getenv("MILVUS_COLLECTION", "kb_content_units_v1")
INDEX = os.getenv("OPENSEARCH_INDEX", "kb-content-units-v1")
LOCAL_DIM = int(os.getenv("LOCAL_EMBEDDING_DIM", "384"))


def value(name: str, mandatory: bool = False) -> str | None:
    result = os.getenv(name) or None
    if mandatory and not result:
        raise RuntimeError(f"Missing setting: {name}")
    return result


def db() -> pymysql.connections.Connection:
    parsed = urlparse(value("MYSQL_DSN", True))
    return pymysql.connect(host=parsed.hostname, port=parsed.port or 3306, user=unquote(parsed.username or ""),
        password=unquote(parsed.password or ""), database=parsed.path.lstrip("/"), charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor, autocommit=False)


def minio() -> Minio:
    return Minio(value("MINIO_ENDPOINT", True), access_key=value("MINIO_ACCESS_KEY", True),
        secret_key=value("MINIO_SECRET_KEY", True), secure=False)


def claim() -> dict | None:
    with db() as conn, conn.cursor() as cur:
        cur.execute("SELECT j.id job_id,j.job_type,j.document_version_id,v.object_key,v.original_filename,"
            "d.id document_id,d.knowledge_base_id FROM ingestion_job j "
            "JOIN document_version v ON v.id=j.document_version_id JOIN document d ON d.id=v.document_id "
            "WHERE j.status='queued' AND j.job_type IN ('extract','reindex','delete') "
            "ORDER BY j.created_at LIMIT 1 FOR UPDATE SKIP LOCKED")
        job = cur.fetchone()
        if not job:
            return None
        cur.execute("UPDATE ingestion_job SET status='running',attempt_count=attempt_count+1,started_at=NOW(3) WHERE id=%s", (job["job_id"],))
        if job["job_type"] != "delete":
            cur.execute("UPDATE document_version SET extraction_status='processing' WHERE id=%s", (job["document_version_id"],))
        conn.commit()
        return job


def finish(job: dict, ok: bool, error: str | None = None) -> None:
    with db() as conn, conn.cursor() as cur:
        if ok:
            cur.execute("UPDATE ingestion_job SET status='succeeded',finished_at=NOW(3),error_message=NULL WHERE id=%s", (job["job_id"],))
            if job["job_type"] != "delete":
                cur.execute("UPDATE document_version SET extraction_status='succeeded',extracted_at=NOW(3),"
                    "parser_name='builtin',parser_version='0.5.0' WHERE id=%s", (job["document_version_id"],))
        else:
            cur.execute("UPDATE ingestion_job SET status='failed',finished_at=NOW(3),error_message=%s WHERE id=%s", ((error or "")[:4000], job["job_id"]))
            if job["job_type"] != "delete":
                cur.execute("UPDATE document_version SET extraction_status='failed' WHERE id=%s", (job["document_version_id"],))
        conn.commit()


def download(job: dict) -> Path:
    descriptor, filename = tempfile.mkstemp(prefix="kb-source-")
    os.close(descriptor)
    path = Path(filename)
    response = minio().get_object(value("MINIO_BUCKET", True), job["object_key"])
    try:
        with path.open("wb") as output:
            for block in response.stream(1024 * 1024):
                output.write(block)
    finally:
        response.close()
        response.release_conn()
    return path


def extract(path: Path, filename: str) -> list[tuple[int | None, str]]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        pages = [(number, (page.extract_text() or "").strip()) for number, page in enumerate(PdfReader(str(path)).pages, 1)]
        if sum(len(text) for _, text in pages) >= 100:
            return pages
        return [(number, pytesseract.image_to_string(image, lang="chi_sim+eng").strip())
                for number, image in enumerate(convert_from_path(str(path), dpi=200), 1)]
    if extension == ".docx":
        document = DocxDocument(str(path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            lines.extend(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells) for row in table.rows)
        return [(None, "\n".join(lines))]
    if extension in {".xlsx", ".xlsm"}:
        workbook, lines = load_workbook(str(path), read_only=True, data_only=True), []
        for sheet in workbook.worksheets:
            lines.append(f"# 工作表：{sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        workbook.close()
        return [(None, "\n".join(lines))]
    if extension in {".txt", ".md", ".csv"}:
        return [(None, path.read_text(encoding="utf-8", errors="ignore"))]
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return [(1, pytesseract.image_to_string(str(path), lang="chi_sim+eng").strip())]
    raise ValueError(f"暂不支持的文件类型：{extension}")


def split_pages(pages: list[tuple[int | None, str]], size: int = 1200, overlap: int = 150) -> list[tuple[int | None, str]]:
    output = []
    for page, text in pages:
        buffer = ""
        for line in (line.strip() for line in text.splitlines() if line.strip()):
            if buffer and len(buffer) + len(line) + 1 > size:
                output.append((page, buffer))
                buffer = buffer[-overlap:] + "\n" + line
            else:
                buffer = (buffer + "\n" + line).strip()
        if buffer:
            output.append((page, buffer))
    return output


def cleanup_indexes(document_id: int) -> None:
    try:
        connections.connect(alias="default", uri=value("MILVUS_URI", True))
        if utility.has_collection(COLLECTION):
            collection = Collection(COLLECTION)
            collection.delete(f"document_id == {document_id}")
            collection.flush()
    except Exception as exc:
        log.warning("清理 Milvus 文档 %s 失败：%s", document_id, exc)
    try:
        response = httpx.post(f"{value('OPENSEARCH_URL', True).rstrip('/')}/{INDEX}/_delete_by_query?refresh=true",
            auth=(value("OPENSEARCH_USERNAME", True), value("OPENSEARCH_PASSWORD", True)), verify=False,
            json={"query": {"term": {"document_id": document_id}}}, timeout=30)
        if response.status_code != 404:
            response.raise_for_status()
    except Exception as exc:
        log.warning("清理 OpenSearch 文档 %s 失败：%s", document_id, exc)


def save_units(job: dict, chunks: list[tuple[int | None, str]]) -> list[dict]:
    cleanup_indexes(job["document_id"])
    with db() as conn, conn.cursor() as cur:
        cur.execute("DELETE FROM content_unit WHERE document_version_id=%s", (job["document_version_id"],))
        units = []
        for sequence, (page, text) in enumerate(chunks, 1):
            cur.execute("INSERT INTO content_unit (document_version_id,unit_type,sequence_no,page_start,page_end,content_text,content_hash,token_count,metadata_json) "
                "VALUES (%s,'chunk',%s,%s,%s,%s,%s,%s,JSON_OBJECT('source','worker'))",
                (job["document_version_id"], sequence, page, page, text, hashlib.sha256(text.encode()).hexdigest(), max(1, len(text) // 3)))
            units.append({"id": cur.lastrowid, "page": page, "text": text})
        conn.commit()
        return units


def local_hash_embedding(text: str) -> list[float]:
    base = re.findall(r"[a-z0-9]+|[\u4e00-\u9fff]", text.lower())
    tokens = base + ["".join(base[index:index + 2]) for index in range(max(0, len(base) - 1))]
    vector = [0.0] * LOCAL_DIM
    for token in tokens:
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % LOCAL_DIM
        vector[index] += 1.0 if digest[4] & 1 else -1.0
    norm = math.sqrt(sum(item * item for item in vector)) or 1.0
    return [item / norm for item in vector]


def embed(text: str) -> list[float] | None:
    if (value("EMBEDDING_PROVIDER") or "local_hash") == "local_hash":
        return local_hash_embedding(text)
    base, model = value("EMBEDDING_BASE_URL"), value("EMBEDDING_MODEL")
    if not base or not model:
        return None
    headers = {"Content-Type": "application/json"}
    if value("EMBEDDING_API_KEY"):
        headers["Authorization"] = f"Bearer {value('EMBEDDING_API_KEY')}"
    response = httpx.post(base.rstrip("/") + "/embeddings", headers=headers,
        json={"model": model, "input": text}, timeout=90)
    response.raise_for_status()
    return response.json()["data"][0]["embedding"]


def milvus_collection(dimension: int) -> Collection:
    connections.connect(alias="default", uri=value("MILVUS_URI", True))
    if not utility.has_collection(COLLECTION):
        schema = CollectionSchema([
            FieldSchema("id", DataType.VARCHAR, is_primary=True, auto_id=False, max_length=64),
            FieldSchema("vector", DataType.FLOAT_VECTOR, dim=dimension),
            FieldSchema("knowledge_base_id", DataType.INT64), FieldSchema("document_id", DataType.INT64),
            FieldSchema("content_unit_id", DataType.INT64)])
        collection = Collection(COLLECTION, schema)
        collection.create_index("vector", {"index_type": "AUTOINDEX", "metric_type": "COSINE", "params": {}})
    else:
        collection = Collection(COLLECTION)
        vector_field = next(field for field in collection.schema.fields if field.name == "vector")
        if vector_field.params.get("dim") != dimension:
            raise RuntimeError(f"Milvus 集合维度为 {vector_field.params.get('dim')}，当前向量维度为 {dimension}")
    collection.load()
    return collection


def ensure_fulltext_index() -> None:
    mapping = {"mappings": {"properties": {"knowledge_base_id": {"type": "long"},
        "document_id": {"type": "long"}, "content_unit_id": {"type": "long"},
        "page": {"type": "integer"}, "text": {"type": "text"}, "department_ids": {"type": "long"}}}}
    response = httpx.put(f"{value('OPENSEARCH_URL', True).rstrip('/')}/{INDEX}",
        auth=(value("OPENSEARCH_USERNAME", True), value("OPENSEARCH_PASSWORD", True)), verify=False,
        json=mapping, timeout=30)
    if response.status_code not in {200, 201, 400}:
        response.raise_for_status()


def index_units(job: dict, units: list[dict]) -> None:
    with db() as conn, conn.cursor() as cur:
        cur.execute("SELECT department_id FROM document_department_acl WHERE document_id=%s", (job["document_id"],))
        department_ids = [row["department_id"] for row in cur.fetchall()]
    ensure_fulltext_index()
    base = value("OPENSEARCH_URL", True).rstrip("/")
    auth = (value("OPENSEARCH_USERNAME", True), value("OPENSEARCH_PASSWORD", True))
    for unit in units:
        payload = {"knowledge_base_id": job["knowledge_base_id"], "document_id": job["document_id"],
            "content_unit_id": unit["id"], "page": unit["page"], "text": unit["text"],
            "department_ids": department_ids}
        response = httpx.put(f"{base}/{INDEX}/_doc/cu-{unit['id']}?refresh=wait_for", auth=auth,
            verify=False, json=payload, timeout=30)
        response.raise_for_status()
    with db() as conn, conn.cursor() as cur:
        cur.execute("UPDATE content_unit SET fulltext_status='indexed',opensearch_id=CONCAT('cu-',id) WHERE document_version_id=%s", (job["document_version_id"],))
        conn.commit()
    vectors = [embed(unit["text"]) for unit in units]
    if not vectors or vectors[0] is None:
        log.warning("Embedding 未配置：全文索引已完成，向量索引保持 pending")
        return
    if any(vector is None for vector in vectors):
        raise RuntimeError("Embedding 服务返回空向量")
    collection = milvus_collection(len(vectors[0]))
    for unit, vector in zip(units, vectors, strict=True):
        vector_id = f"cu-{unit['id']}"
        collection.upsert([[vector_id], [vector], [job["knowledge_base_id"]], [job["document_id"]], [unit["id"]]])
    collection.flush()
    with db() as conn, conn.cursor() as cur:
        cur.execute("UPDATE content_unit SET vector_status='indexed',milvus_pk=CONCAT('cu-',id) WHERE document_version_id=%s", (job["document_version_id"],))
        conn.commit()


def delete_document(job: dict) -> None:
    cleanup_indexes(job["document_id"])
    try:
        minio().remove_object(value("MINIO_BUCKET", True), job["object_key"])
    except Exception as exc:
        log.warning("删除 MinIO 原文件失败：%s", exc)
    with db() as conn, conn.cursor() as cur:
        cur.execute("DELETE FROM content_unit WHERE document_version_id=%s", (job["document_version_id"],))
        conn.commit()


def run(job: dict) -> None:
    if job["job_type"] == "delete":
        delete_document(job)
        finish(job, True)
        log.info("完成删除任务 %s", job["job_id"])
        return
    source = download(job)
    try:
        units = save_units(job, split_pages(extract(source, job["original_filename"])))
        if not units:
            raise ValueError("未提取到可用文本")
        index_units(job, units)
        finish(job, True)
        log.info("完成任务 %s，切片数 %s", job["job_id"], len(units))
    finally:
        source.unlink(missing_ok=True)


def main() -> None:
    log.info("knowledge-base worker started")
    while True:
        job = None
        try:
            job = claim()
            if not job:
                time.sleep(3)
            else:
                run(job)
        except Exception as exc:
            log.exception("任务处理失败")
            if job:
                finish(job, False, str(exc))
            time.sleep(2)


if __name__ == "__main__":
    main()
