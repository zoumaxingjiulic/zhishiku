"""Structure-preserving extraction. Heuristics are conservative, not layout AI.

No database/network dependencies: extraction and chunking can be regression-tested
without touching uploaded documents or calling an embedding service.
"""
from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
import re

PARSER_VERSION = "0.8.0"
HEADING = re.compile(r"^(?:第[一二三四五六七八九十百零\d]+[章节条]|[一二三四五六七八九十]+[、．]|\d+(?:\.\d+)*[、．\s]|#{1,6}\s)")
PAGE_NUMBER = re.compile(r"^(?:[-—–]\s*)?(?:第\s*)?\d+\s*(?:页(?:\s*[/／共]\s*\d+\s*页?)?|[/／]\s*\d+)?(?:\s*[-—–])?$|^Page\s+\d+(?:\s+of\s+\d+)?$", re.I)


@dataclass
class Block:
    text: str
    kind: str = "text"
    page_start: int | None = None
    page_end: int | None = None
    context: str = ""
    metadata: dict = field(default_factory=dict)


@dataclass
class Chunk:
    text: str
    page_start: int | None
    page_end: int | None
    metadata: dict


def clean(value) -> str:
    return re.sub(r"[ \t]+", " ", str(value) if value is not None else "").strip()


def join_text(left: str, right: str) -> str:
    # Chinese needs no injected space. Keep Latin words separated; do not guess
    # whether a trailing hyphen is a word wrap or a meaningful product number.
    gap = " " if left and right and left[-1].isascii() and right[0].isascii() else ""
    return left + gap + right


def header_candidate(row: list[str]) -> bool:
    present = [v for v in row if v]
    return len(present) >= 2 and all(not re.fullmatch(r"[-+\d.,%/年月日 :]+", v) for v in present)


def table_blocks(rows: list[tuple[int, list[str]]], *, context: str, metadata: dict,
                 page: int | None = None, headers: list[str] | None = None) -> list[Block]:
    if not rows:
        return []
    width = max(len(row) for _, row in rows)
    header_row = None
    if headers is None and header_candidate(rows[0][1]):
        header_row, headers = rows[0]
        rows = rows[1:]
    labels = [(headers[i] if headers and i < len(headers) and headers[i] else f"列{i + 1}") for i in range(width)]
    prefix = context + "\n表头：" + " | ".join(labels)
    common = {**metadata, "headers": labels, "header_row": metadata.get("header_row", header_row)}
    if not rows:
        return [Block("（表格无数据行）", "table", page, page, prefix, common)]
    return [Block(f"第{number}行：" + "；".join(f"{label}：{row[i] if i < len(row) else ''}" for i, label in enumerate(labels)),
                  "table", page, page, prefix, {**common, "row_start": number, "row_end": number})
            for number, row in rows]


def extract_docx(path: Path) -> list[Block]:
    from docx import Document
    from docx.table import Table

    document, result, headings = Document(str(path)), [], []
    table_number = 0
    # iter_inner_content preserves body paragraph/table order, unlike collecting
    # document.paragraphs and document.tables in separate passes.
    for item in document.iter_inner_content():
        if isinstance(item, Table):
            table_number += 1
            rows = [(n, [clean(cell.text) for cell in row.cells]) for n, row in enumerate(item.rows, 1)]
            result.extend(table_blocks(rows, context=" / ".join(headings) + f"\n表格 {table_number}",
                                       metadata={"format": "docx", "table_id": table_number, "headings": headings[:]}))
            continue
        text = item.text.strip()
        if not text:
            continue
        style = item.style.name if item.style else ""
        match = re.search(r"(?:Heading|标题)\s*(\d+)", style, re.I)
        outline = item._p.xpath("./w:pPr/w:outlineLvl")
        level = int(match[1]) if match else None
        if outline:
            from docx.oxml.ns import qn
            outline_level = int(outline[0].get(qn("w:val")))
            if outline_level < 9:
                level = outline_level + 1
        if level:
            headings = headings[:level - 1] + [text]
        result.append(Block(text, "heading" if level else "text", context=" / ".join(headings),
                            metadata={"format": "docx", "headings": headings[:], "style": style}))
    return result


def extract_excel(path: Path) -> list[Block]:
    from openpyxl import load_workbook

    result = []
    # Normal mode is needed for merged header cells. Values are cached formula
    # results: openpyxl must not evaluate formulas or execute workbook macros.
    # Worker downloads objects to extensionless mkstemp paths. Passing a filename
    # makes openpyxl reject those; a binary stream uses the actual ZIP content.
    with path.open("rb") as source:
        workbook = load_workbook(source, data_only=True)
    try:
        for sheet in workbook.worksheets:
            regions, region = [], []
            for cells in sheet.iter_rows():
                row = [clean(cell.value) if cell.value is not None else "" for cell in cells]
                if any(row):
                    region.append((cells[0].row, row))
                elif region:
                    regions.append(region)
                    region = []
            if region:
                regions.append(region)
            for table_id, rows in enumerate(regions, 1):
                # Ignore styled-but-empty far-right columns.
                width = max(max(i + 1 for i, value in enumerate(row) if value) for _, row in rows)
                relevant_merges = [m for m in sheet.merged_cells.ranges if rows[0][0] <= m.min_row <= rows[-1][0]]
                width = max([width] + [m.max_col for m in relevant_merges if m.min_row <= rows[0][0] + 2])
                rows = [(n, row[:width] + [""] * max(0, width - len(row))) for n, row in rows]
                titles = []
                # A single-cell leading row is treated as a caption only when
                # another multi-column row follows; its content is never dropped.
                while len(rows) > 1 and sum(bool(v) for v in rows[0][1]) == 1 and width > 1:
                    titles.append(next(v for v in rows.pop(0)[1] if v))
                context = f"工作表：{sheet.title}" + ("\n表名：" + " / ".join(titles) if titles else "")
                metadata = {"format": "xlsx", "sheet": sheet.title, "table_id": table_id, "titles": titles}
                headers = None
                # Explicit merged header geometry, followed by another header
                # row, allows a bounded multi-level header. Do not forward-fill
                # merged body cells, which could invent data.
                header_rows = []
                while rows and len(header_rows) < 3:
                    number, values = rows[0]
                    expanded = values[:]
                    merges = [m for m in relevant_merges if m.min_row <= number <= m.max_row]
                    for merged in merges:
                        for col in range(merged.min_col - 1, min(merged.max_col, width)):
                            expanded[col] = clean(sheet.cell(merged.min_row, merged.min_col).value)
                    if not header_candidate(expanded):
                        break
                    header_rows.append((number, expanded))
                    rows.pop(0)
                    if not merges or not rows:
                        break
                if header_rows:
                    headers = [" / ".join(dict.fromkeys(row[i] for _, row in header_rows if row[i])) for i in range(width)]
                    metadata["header_row"] = header_rows[0][0]
                    metadata["header_row_end"] = header_rows[-1][0]
                result.extend(table_blocks(rows, context=context, metadata=metadata, headers=headers) if rows else
                              [Block("（表格无数据行）", "table", context=context + "\n表头：" + " | ".join(headers or []), metadata=metadata)])
    finally:
        workbook.close()
    return result


@dataclass
class PdfPage:
    number: int
    width: float
    height: float
    elements: list[dict]


def _margin_key(text: str) -> str:
    text = clean(text)
    return "<page-number>" if PAGE_NUMBER.fullmatch(text) else text


def clean_pdf_margins(pages: list[PdfPage]) -> None:
    """Only repeated text in narrow top/bottom bands; never global digit stripping."""
    counts = Counter()
    for page in pages:
        keys = set()
        for el in page.elements:
            if el["kind"] != "text":
                continue
            band = "top" if el["bottom"] < page.height * .08 else "bottom" if el["top"] > page.height * .92 else None
            if band:
                keys.add((band, _margin_key(el["text"])))
        counts.update(keys)
    required = max(2, (len(pages) + 1) // 2)
    for page in pages:
        kept = []
        for el in page.elements:
            band = "top" if el["bottom"] < page.height * .08 else "bottom" if el["top"] > page.height * .92 else None
            key = _margin_key(el.get("text", ""))
            if el["kind"] == "text" and band and (key == "<page-number>" or counts[(band, key)] >= required):
                continue
            kept.append(el)
        page.elements = kept


def _continuation(previous: dict, current: dict, before: PdfPage, after: PdfPage) -> bool:
    if after.number != before.number + 1 or previous["kind"] != current["kind"]:
        return False
    if previous["bottom"] < before.height * .78 or current["top"] > after.height * .22:
        return False
    if abs(previous["x0"] / before.width - current["x0"] / after.width) > .04:
        return False
    if previous["kind"] == "table":
        # Same geometry alone is insufficient: unrelated tables can have the
        # same number of columns. Require an identical, non-empty header too.
        return (len(previous["rows"]) > 1 and len(current["rows"]) > 1
                and header_candidate(previous["rows"][0])
                and previous["rows"][0] == current["rows"][0]
                and abs(previous["x1"] / before.width - current["x1"] / after.width) < .04)
    left, right = previous["text"], current["text"]
    return (not previous.get("heading") and not current.get("heading") and not HEADING.match(right)
            and not re.search(r"[。！？.!?；;：:]\s*$", left)
            and not re.match(r"[（(]?[一二三四五六七八九十\d]+[）)、.]", right))


def assemble_pdf(pages: list[PdfPage]) -> list[Block]:
    clean_pdf_margins(pages)
    result, previous_page, previous_element = [], None, None
    table_id, headings = 0, []
    for page in pages:
        for index, el in enumerate(page.elements):
            continued = bool(index == 0 and previous_element and previous_page and
                             _continuation(previous_element, el, previous_page, page))
            if el["kind"] == "table":
                if not continued:
                    table_id += 1
                rows = el["rows"]
                offset = int(result[-1].metadata.get("row_end", 1)) if continued and result else 0
                numbered = [(n + offset, row) for n, row in enumerate(rows[1:] if continued else rows, 1)]
                blocks = table_blocks(numbered, context=" / ".join(headings) + f"\n表格 {table_id}",
                    metadata={"format": "pdf", "table_id": table_id, "continued_table": continued,
                              "extraction": el.get("extraction", "digital"), "headings": headings[:]},
                    page=page.number, headers=rows[0] if continued else None)
                result.extend(blocks)
            else:
                text = el["text"]
                heading = el.get("heading") or bool(HEADING.match(text) and len(text) < 100)
                if heading:
                    headings = [text]
                metadata = {"format": "pdf", "headings": headings[:], "extraction": el.get("extraction", "digital")}
                if continued and result:
                    # Merge only the actual continuation, retaining both pages.
                    result[-1].text = join_text(result[-1].text, text)
                    result[-1].page_end = page.number
                    result[-1].metadata["cross_page_continuation"] = True
                else:
                    result.append(Block(text, "heading" if heading else "text", page.number, page.number,
                                        " / ".join(headings), metadata))
            previous_element = el
        previous_page = page
        if not page.elements:
            previous_element = None  # never join across blank/unreadable pages
    return result


def _ocr_pdf_page(path: Path, number: int) -> PdfPage:
    import pytesseract
    from pdf2image import convert_from_path

    images = convert_from_path(str(path), dpi=200, first_page=number, last_page=number)
    image = images[0]
    try:
        data = pytesseract.image_to_data(image, lang="chi_sim+eng", output_type=pytesseract.Output.DICT)
        lines = {}
        for i, text in enumerate(data["text"]):
            if not text.strip():
                continue
            key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
            lines.setdefault(key, []).append((text, data["left"][i], data["top"][i], data["width"][i], data["height"][i]))
        elements = [{"kind": "text", "text": " ".join(v[0] for v in words),
                     "x0": min(v[1] for v in words), "x1": max(v[1] + v[3] for v in words),
                     "top": min(v[2] for v in words), "bottom": max(v[2] + v[4] for v in words), "extraction": "ocr"}
                    for words in lines.values()]
        return PdfPage(number, image.width, image.height, sorted(elements, key=lambda e: (e["top"], e["x0"])))
    finally:
        for image in images:
            image.close()


def extract_pdf(path: Path) -> list[Block]:
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as document:
        for number, page in enumerate(document.pages, 1):
            # A scanned body may still have a long native-text header/footer.
            # Inspect body text and image area, not document-wide text length.
            body_chars = [c for c in page.chars if page.height * .08 <= c["top"] <= page.height * .92]
            large_image = any((image["x1"] - image["x0"]) * (image["bottom"] - image["top"])
                              >= page.width * page.height * .2 for image in page.images)
            if len(body_chars) < 30 and large_image:
                pages.append(_ocr_pdf_page(path, number))
                page.close()
                continue
            tables = [t for t in page.find_tables() if len(t.rows) >= 2 and len(t.columns) >= 2]
            elements = []
            for table in tables:
                rows = [[clean(cell) for cell in row] for row in table.extract()]
                if rows:
                    elements.append({"kind": "table", "rows": rows, "x0": table.bbox[0], "top": table.bbox[1],
                                     "x1": table.bbox[2], "bottom": table.bbox[3]})
            def outside_table(obj):
                return not any(t.bbox[0] <= (obj["x0"] + obj["x1"]) / 2 <= t.bbox[2]
                               and t.bbox[1] <= (obj["top"] + obj["bottom"]) / 2 <= t.bbox[3] for t in tables)
            body = page.filter(outside_table)
            typical_size = sorted(c["size"] for c in page.chars)[len(page.chars) // 2] if page.chars else 12
            for line in body.extract_text_lines():
                text = line["text"].strip()
                if text:
                    sizes = [c["size"] for c in line.get("chars", [])]
                    elements.append({**{key: line[key] for key in ("x0", "x1", "top", "bottom")},
                                     "kind": "text", "text": text,
                                     "heading": bool(sizes and min(sizes) > typical_size * 1.2)})
            pages.append(PdfPage(number, page.width, page.height, sorted(elements, key=lambda e: (e["top"], e["x0"]))))
            page.close()
    return assemble_pdf(pages)


def extract(path: Path, filename: str) -> list[Block]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".docx":
        return extract_docx(path)
    if extension in {".xlsx", ".xlsm"}:
        return extract_excel(path)
    if extension in {".txt", ".md", ".csv"}:
        return [Block(path.read_text(encoding="utf-8", errors="replace"), metadata={"format": extension[1:]})]
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        import pytesseract
        return [Block(pytesseract.image_to_string(str(path), lang="chi_sim+eng").strip(),
                      page_start=1, page_end=1, metadata={"format": "image", "extraction": "ocr"})]
    raise ValueError(f"暂不支持的文件类型：{extension}")


def split_blocks(blocks: list[Block], size: int = 1200, overlap: int = 150) -> list[Chunk]:
    if size < 64 or overlap < 0 or overlap >= size:
        raise ValueError("切片长度须 >= 64，重叠长度须 >= 0 且小于切片长度")
    output, group = [], []

    def flush():
        if not group:
            return
        context = group[0].context
        # Never silently truncate headers: stop this ingestion with an actionable
        # error if a table is too wide to carry its context into every chunk.
        if len(context) >= size - 32:
            raise ValueError("表头或标题上下文过长，请拆分宽表/缩短标题后重新上传")
        prefix = context + "\n" if context else ""
        body, spans = "", []
        for block in group:
            if body:
                body += "\n"
            start = len(body)
            body += block.text
            spans.append((start, len(body), block))
        start, budget = 0, size - len(prefix)
        effective_overlap = min(overlap, budget // 4)
        while start < len(body):
            end = min(start + budget, len(body))
            if end < len(body):
                floor = start + budget // 2
                boundaries = [body.rfind(sep, floor, end) for sep in ("\n", "。", "；", ". ", " ")]
                boundary = max(boundaries)
                if boundary >= floor:
                    end = boundary + 1
            selected = [block for left, right, block in spans if left < end and right > start]
            pages = [p for b in selected for p in (b.page_start, b.page_end) if p is not None]
            metadata = {**selected[0].metadata, "source": "worker", "parser_version": PARSER_VERSION,
                        "kind": selected[0].kind, "chunk_size": size, "overlap": effective_overlap,
                        "sources": [{"page_start": b.page_start, "page_end": b.page_end, **b.metadata} for b in selected]}
            row_numbers = [b.metadata["row_start"] for b in selected if "row_start" in b.metadata]
            if row_numbers:
                metadata.update(row_start=min(row_numbers), row_end=max(row_numbers))
            text = prefix + body[start:end].strip()
            if text.strip():
                output.append(Chunk(text, min(pages) if pages else None, max(pages) if pages else None, metadata))
            if end == len(body):
                break
            start = end - effective_overlap
        group.clear()

    for block in blocks:
        if not block.text.strip():
            continue
        # Keep different worksheets/tables/sections apart. Across PDF pages only
        # assemble_pdf's positively identified continuation may span the boundary.
        key = (block.kind == "table", block.context, block.metadata.get("table_id"), block.page_start)
        if group:
            previous = group[-1]
            old_key = (previous.kind == "table", previous.context, previous.metadata.get("table_id"), previous.page_start)
            same_continued_table = (block.kind == previous.kind == "table" and block.context == previous.context
                                    and block.metadata.get("continued_table") and block.metadata.get("table_id") == previous.metadata.get("table_id"))
            same_continued_paragraph = (previous.kind == block.kind == "text" and previous.context == block.context
                                       and previous.metadata.get("cross_page_continuation")
                                       and previous.page_end == block.page_start)
            if key != old_key and not same_continued_table and not same_continued_paragraph:
                flush()
        group.append(block)
    flush()
    return output
